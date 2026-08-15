#!/usr/bin/env python3
"""Formatiert die fertige Architektur-Seminararbeit im Layout der Vorlage
``Seminararbeit_Steering_Committee.docx``.

Die Vorlage gibt vor:
  * Fließtext Times New Roman 12 pt, Blocksatz, Zeilenabstand 1,3.
  * Überschriften Times New Roman, fett, schwarz (H1 14 pt, H2 13 pt, H3 12 pt),
    ohne farbige Linien.
  * Zentriertes, ruhiges Deckblatt (Hochschule / Titel / Untertitel / Modul /
    Studiengang / Semester / Vorgelegt von / Prüfer / Abgabedatum).
  * Seitenränder 3,0 cm (oben/links) und 2,5 cm (unten/rechts), A4.
  * Schlichte Tabellen (fette Kopfzeile, Gitternetz, keine Flächenfüllung).

Abbildungen werden proportional auf die nutzbare Seitenbreite/-höhe begrenzt.
Der Autoren-Kopf je Kapitel (Gruppenarbeit) bleibt erhalten, an die Serifen-
Vorlage angepasst.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

SERIF = "Times New Roman"
GREY = RGBColor(0x66, 0x66, 0x66)
BLACK = "000000"

# Nutzbare Seitenbreite/-höhe (A4 21,0 x 29,7 cm, Ränder 3,0/2,5 cm).
CONTENT_WIDTH_CM = 21.0 - 3.0 - 2.5          # 15,5 cm
CONTENT_HEIGHT_CM = 29.7 - 3.0 - 2.5 - 1.6   # abzüglich Beschriftung/Reserve

COVER = {
    "hochschule": "Duale Hochschule Baden-Württemberg\nCenter for Advanced Studies",
    "title": "Modernisierung der Vereinsverwaltung FH_MA \u2014 "
    "Architekturkonzeption und Umsetzungsevaluation",
    "subtitle": "Seminararbeit der Arbeitsgruppe 3 (Architektur)",
    "modul": "Seminararbeit im Rahmen des Moduls\n"
    "Advanced Software Engineering (CSC1200)",
    "studiengang": "Studiengang: M.Sc. Informatik",
    "semester": "Semester: Sommersemester 2026",
    "authors": [
        "Vorgelegt von:",
        "Nils Firschau (8993076)  \u00b7  Paul Faller (5567855)",
        "Robin Steiner (9251426)  \u00b7  Ole Schildt (3504736)",
    ],
    "pruefer": "Pr\u00fcfer: Prof. Dr. Holger D. Hofmann",
    "abgabe": "Abgabedatum: 15. August 2026",
}


def set_style(document, name, font, size, *, bold=False, color=None, alignment=None):
    style = document.styles[name]
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    # Theme-Verweise entfernen, damit die explizite Schrift greift (LibreOffice
    # bevorzugt sonst die Theme-Schrift majorHAnsi = serifenlos).
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if alignment is not None:
        style.paragraph_format.alignment = alignment
    return style


def clear_paragraph_borders(paragraph):
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        for bdr in p_pr.findall(qn("w:pBdr")):
            p_pr.remove(bdr)


def make_run(text, *, font=SERIF, size=None, bold=False, color=None):
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    rpr.append(rfonts)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rpr.append(sz)
    if color:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), color)
        rpr.append(col)
    run.append(rpr)
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i:
            run.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = part
        run.append(t)
    return run


def make_paragraph(text, *, size, bold=False, space_after=6, space_before=0):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(int(space_after * 20)))
    spacing.set(qn("w:before"), str(int(space_before * 20)))
    p_pr.append(spacing)
    p.append(p_pr)
    if text:
        p.append(make_run(text, size=size, bold=bold, color=BLACK))
    return p


def page_break_paragraph():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def rebuild_cover(document):
    """Entfernt das alte Deckblatt und setzt das Layout der Vorlage ein."""
    body = document.element.body
    # Erste Inhaltsüberschrift = Grenze des Deckblatts.
    boundary = None
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            style = child.find(qn("w:pPr"))
            pstyle = None
            if style is not None:
                ps = style.find(qn("w:pStyle"))
                if ps is not None:
                    pstyle = ps.get(qn("w:val"))
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if pstyle in {"FrontHeading"} or text.strip() == "Inhaltsverzeichnis":
                boundary = child
                break
    if boundary is None:
        raise RuntimeError("Deckblattgrenze (Inhaltsverzeichnis) nicht gefunden.")

    # Alte Deckblatt-Elemente entfernen.
    for child in list(body.iterchildren()):
        if child is boundary:
            break
        if child.tag in {qn("w:p"), qn("w:tbl")}:
            body.remove(child)

    blocks = [
        ("spacer", 1),
        (COVER["hochschule"], dict(size=13, bold=True, space_after=6)),
        ("spacer", 3),
        (COVER["title"], dict(size=20, bold=True, space_after=6)),
        (COVER["subtitle"], dict(size=14, bold=False, space_after=6)),
        ("spacer", 3),
        (COVER["modul"], dict(size=14, bold=False, space_after=6)),
        ("spacer", 1),
        (COVER["studiengang"], dict(size=12, bold=False, space_after=2)),
        (COVER["semester"], dict(size=12, bold=False, space_after=6)),
        ("spacer", 2),
    ]
    for line in COVER["authors"]:
        blocks.append((line, dict(size=12, bold=False, space_after=2)))
    blocks += [
        ("spacer", 1),
        (COVER["pruefer"], dict(size=12, bold=False, space_after=6)),
        ("spacer", 1),
        (COVER["abgabe"], dict(size=12, bold=False, space_after=6)),
    ]

    for item, cfg in blocks:
        if item == "spacer":
            for _ in range(cfg):
                boundary.addprevious(make_paragraph("", size=12, space_after=6))
        else:
            boundary.addprevious(make_paragraph(item, **cfg))

    # Inhaltsverzeichnis auf neue Seite zwingen (ohne separate Leerseite).
    b_pr = boundary.find(qn("w:pPr"))
    if b_pr is None:
        b_pr = OxmlElement("w:pPr")
        boundary.insert(0, b_pr)
    if b_pr.find(qn("w:pageBreakBefore")) is None:
        b_pr.insert(0, OxmlElement("w:pageBreakBefore"))


def restyle_images(document):
    """Begrenzt jede Abbildung proportional auf die nutzbare Seitenfläche."""
    max_w = Cm(CONTENT_WIDTH_CM)
    max_h = Cm(CONTENT_HEIGHT_CM)
    for inline in document.element.body.iter(qn("wp:inline")):
        extent = inline.find(qn("wp:extent"))
        cx = int(extent.get("cx"))
        cy = int(extent.get("cy"))
        scale = min(max_w / cx, max_h / cy, 1.0)
        # kleine Sicherheitsreserve zur Breite
        scale = min(scale, (Cm(CONTENT_WIDTH_CM - 0.4)) / cx)
        new_cx = int(cx * scale)
        new_cy = int(cy * scale)
        extent.set("cx", str(new_cx))
        extent.set("cy", str(new_cy))
        for ext in inline.iter(qn("a:ext")):
            ext.set("cx", str(new_cx))
            ext.set("cy", str(new_cy))


def style_header(section, author):
    header = section.header
    paragraph = header.paragraphs[0]
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    for run in paragraph.runs:
        pass
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for tab in list(paragraph.paragraph_format.tab_stops):
        pass
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.05), WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run(f"Arbeitsgruppe 3 \u00b7 Architektur\t{author}")
    run.font.name = SERIF
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    clear_paragraph_borders(paragraph)


def add_page_field(paragraph):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), SERIF)
    rfonts.set(qn("w:hAnsi"), SERIF)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def set_page_numbering(section, fmt, start=None):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is None:
        existing = OxmlElement("w:pgNumType")
        insertion_point = None
        for tag in ("w:cols", "w:formProt", "w:vAlign", "w:titlePg", "w:docGrid"):
            insertion_point = sect_pr.find(qn(tag))
            if insertion_point is not None:
                break
        if insertion_point is None:
            sect_pr.append(existing)
        else:
            sect_pr.insert(list(sect_pr).index(insertion_point), existing)
    existing.set(qn("w:fmt"), fmt)
    if start is None:
        existing.attrib.pop(qn("w:start"), None)
    else:
        existing.set(qn("w:start"), str(start))


def insert_section_break_before(paragraph):
    preceding = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "nextPage")
    sect_pr.append(section_type)
    p_pr.append(sect_pr)
    preceding.append(p_pr)
    paragraph._p.addprevious(preceding)


def format_document(source: Path, target: Path, page_map: dict | None = None):
    document = Document(source)

    # --- Seitenränder / Abstände ---
    for section in document.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        section.gutter = Cm(0)

    # --- Fließtext ---
    normal = set_style(document, "Normal", SERIF, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # --- Überschriften: Serifen, fett, schwarz, keine Linien ---
    for name, size, before in (
        ("Heading 1", 14, 18),
        ("Heading 2", 13, 12),
        ("Heading 3", 12, 10),
    ):
        style = set_style(document, name, SERIF, size, bold=True, color=BLACK)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(before)
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    # Verzeichnis-Überschriften wie H1 (Serif, fett, schwarz, ohne Linie).
    if "FrontHeading" in [s.name for s in document.styles]:
        fh = set_style(
            document, "FrontHeading", SERIF, 14, bold=True, color=BLACK,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        fh.paragraph_format.space_before = Pt(0)
        fh.paragraph_format.space_after = Pt(10)
        fh.paragraph_format.keep_with_next = True

    # Bild-/Tabellen-/Listing-Beschriftungen: Serif, zentriert.
    for name in ("Abbildungsbeschriftung", "Tabellenbeschriftung", "Listingbeschriftung"):
        if name in [s.name for s in document.styles]:
            cap = set_style(document, name, SERIF, 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            cap.paragraph_format.keep_with_next = True
            cap.paragraph_format.space_after = Pt(8)
            cap.paragraph_format.space_before = Pt(2)

    # Verzeichniseinträge: Serif mit Punktführung.
    for level in range(1, 4):
        name = f"DirectoryEntry{level}"
        if name in [s.name for s in document.styles]:
            de = set_style(document, name, SERIF, 11.5 - (level - 1) * 0.5,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
            de.paragraph_format.left_indent = Cm((level - 1) * 0.5)
            de.paragraph_format.space_after = Pt(2)
            de.paragraph_format.line_spacing = 1.15
            de.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            for existing in list(de.paragraph_format.tab_stops):
                pass
            de.paragraph_format.tab_stops.add_tab_stop(
                Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )

    # Cover-Reststile neutralisieren (falls noch referenziert).
    for name in ("WordTitle", "WordSubtitle", "Institution", "InstitutionSub",
                 "DocumentKind", "TitleRule"):
        if name in [s.name for s in document.styles]:
            set_style(document, name, SERIF, 12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if "Source Code" in [s.name for s in document.styles]:
        set_style(document, "Source Code", "Consolas", 8.5)
        document.styles["Source Code"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        document.styles["Source Code"].paragraph_format.space_after = Pt(0)

    # Farbige Linien an Überschriften/Regeln entfernen.
    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3",
                                    "FrontHeading", "TitleRule"}:
            clear_paragraph_borders(paragraph)

    # --- Deckblatt neu aufbauen ---
    rebuild_cover(document)

    # --- Verzeichnis-Seitenzahlen aktualisieren ---
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("DirectoryEntry"):
            target_text = paragraph.text.rsplit("\t", 1)[0].strip()
            target_text = re.sub(r"\s+\?\?$", "", target_text)
            page = page_map.get(target_text, "??") if page_map else "??"
            paragraph.text = f"{target_text}\t{page}"
            paragraph.style = document.styles[paragraph.style.name]

    # --- Tabellen: schlicht (fette Kopfzeile, Gitter, ohne Flächenfüllung) ---
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                for shd in tc_pr.findall(qn("w:shd")):
                    tc_pr.remove(shd)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.1
                    for run in paragraph.runs:
                        run.font.name = SERIF
                        run.font.size = Pt(9.5)
                        if row_index == 0:
                            run.bold = True

    # --- Abbildungen proportional begrenzen ---
    restyle_images(document)

    # --- Abschnittswechsel je Kapitelblock einfügen (Autoren-Kopf) ---
    chapter_markers = {
        "1. Einleitung",
        "4. Implementierte Systemarchitektur",
        "6. Persistenz, Audit und Dokumente",
        "9. Risiken und technische Schulden",
    }
    found = set()
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() in chapter_markers:
            insert_section_break_before(paragraph)
            found.add(paragraph.text.strip())
    missing = chapter_markers - found
    if missing:
        raise RuntimeError(f"Kapitelanfänge nicht gefunden: {sorted(missing)}")

    intermediate = target.with_suffix(".sections.docx")
    document.save(intermediate)
    document = Document(intermediate)
    intermediate.unlink()

    if len(document.sections) < 5:
        raise RuntimeError(
            f"Erwartet wurden 5 Word-Abschnitte, gefunden: {len(document.sections)}"
        )

    # --- Kopf-/Fußzeilen an Vorlage angleichen ---
    authors = ("Nils Firschau", "Nils Firschau", "Paul Faller",
               "Robin Steiner", "Ole Schildt")
    for index, section in enumerate(document.sections):
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        section.gutter = Cm(0)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        if index == 0:
            section.different_first_page_header_footer = True
            section.first_page_header.is_linked_to_previous = False
            section.first_page_header.paragraphs[0].clear()
            section.first_page_footer.is_linked_to_previous = False
            section.first_page_footer.paragraphs[0].clear()
        style_header(section, authors[min(index, len(authors) - 1)])
        add_page_field(section.footer.paragraphs[0])
        if index == 0:
            set_page_numbering(section, "upperRoman", 0)
        elif index == 1:
            set_page_numbering(section, "decimal", 1)
        else:
            set_page_numbering(section, "decimal", None)

    document.save(target)


def main() -> int:
    if len(sys.argv) not in {3, 4}:
        print("Usage: apply_steering_template.py <in.docx> <out.docx> [page-map.json]",
              file=sys.stderr)
        return 2
    page_map = None
    if len(sys.argv) == 4:
        page_map = json.loads(Path(sys.argv[3]).read_text(encoding="utf-8-sig"))
    format_document(Path(sys.argv[1]), Path(sys.argv[2]), page_map)
    print("Vorlagen-Layout angewendet ->", sys.argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
