#!/usr/bin/env python3
"""Formatiert den Pandoc-Word-Export im Stil der Referenzarbeit."""

from __future__ import annotations

import re
import sys
import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "243B53"
LIGHT_GRAY = "E7E6E6"
RED = "C8193C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        insertion_point = None
        for tag in (
            "w:noWrap",
            "w:tcMar",
            "w:textDirection",
            "w:tcFitText",
            "w:vAlign",
            "w:hideMark",
            "w:headers",
        ):
            insertion_point = tc_pr.find(qn(tag))
            if insertion_point is not None:
                break
        if insertion_point is None:
            tc_pr.append(shading)
        else:
            tc_pr.insert(list(tc_pr).index(insertion_point), shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def set_bottom_border(paragraph, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_style(
    document: Document,
    name: str,
    font: str,
    size: float,
    *,
    bold: bool = False,
    color: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    style = document.styles[name]
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if alignment is not None:
        style.paragraph_format.alignment = alignment


def add_simple_field(paragraph, instruction: str, placeholder: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def insert_section_break_before(paragraph) -> None:
    preceding = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "nextPage")
    sect_pr.append(section_type)
    p_pr.append(sect_pr)
    preceding.append(p_pr)
    paragraph._p.addprevious(preceding)


def set_page_numbering(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is None:
        existing = OxmlElement("w:pgNumType")
        insertion_point = None
        for tag in ("w:cols", "w:formProt", "w:vAlign", "w:titlePg", "w:docGrid"):
            insertion_point = sect_pr.find(qn(tag))
            if insertion_point is not None:
                break
        if insertion_point is None:
            sect_pr.append(existing)
        else:
            sect_pr.insert(list(sect_pr).index(insertion_point), existing)
    existing.set(qn("w:fmt"), fmt)
    if start is None:
        existing.attrib.pop(qn("w:start"), None)
    else:
        existing.set(qn("w:start"), str(start))


def set_header(section, author: str, blank_first: bool = False) -> None:
    section.header.is_linked_to_previous = False
    section.different_first_page_header_footer = blank_first
    paragraph = section.header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.25), WD_TAB_ALIGNMENT.RIGHT
    )
    run = paragraph.add_run(f"Arbeitsgruppe 3 · Architektur\t{author}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(102, 102, 102)
    if blank_first:
        section.first_page_header.is_linked_to_previous = False
        section.first_page_header.paragraphs[0].clear()


def set_footer(section, *, fmt: str, start: int | None, blank_first: bool = False) -> None:
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = blank_first
    paragraph = section.footer.paragraphs[0]
    paragraph.clear()
    add_page_field(paragraph)
    set_page_numbering(section, fmt, start)
    if blank_first:
        section.first_page_footer.is_linked_to_previous = False
        section.first_page_footer.paragraphs[0].clear()


def format_document(
    source: Path,
    target: Path,
    page_map: dict[str, str] | None = None,
) -> None:
    document = Document(source)

    for section in document.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.gutter = Cm(0)

    set_style(document, "Normal", "Times New Roman", 11)
    normal = document.styles["Normal"]
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    for name, size, level in (
        ("Heading 1", 16, 1),
        ("Heading 2", 13, 2),
        ("Heading 3", 11.5, 3),
    ):
        set_style(document, name, "Arial", size, bold=True, color=NAVY)
        style = document.styles[name]
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_after = Pt(8 if level == 1 else 5)
        style.paragraph_format.space_before = Pt(0 if level == 1 else 10)
        if level == 1:
            style.paragraph_format.page_break_before = True

    custom = {
        "Institution": ("Arial", 31, True, WD_ALIGN_PARAGRAPH.CENTER),
        "InstitutionSub": ("Arial", 10, False, WD_ALIGN_PARAGRAPH.CENTER),
        "DocumentKind": ("Arial", 14, True, WD_ALIGN_PARAGRAPH.CENTER),
        "WordTitle": ("Arial", 24, True, WD_ALIGN_PARAGRAPH.CENTER),
        "WordSubtitle": ("Times New Roman", 13, False, WD_ALIGN_PARAGRAPH.CENTER),
        "FrontHeading": ("Arial", 16, True, WD_ALIGN_PARAGRAPH.LEFT),
        "Abbildungsbeschriftung": (
            "Times New Roman",
            10,
            False,
            WD_ALIGN_PARAGRAPH.CENTER,
        ),
        "Tabellenbeschriftung": (
            "Times New Roman",
            10,
            False,
            WD_ALIGN_PARAGRAPH.CENTER,
        ),
        "Listingbeschriftung": (
            "Times New Roman",
            10,
            False,
            WD_ALIGN_PARAGRAPH.CENTER,
        ),
        "DirectoryEntry1": ("Times New Roman", 10, False, WD_ALIGN_PARAGRAPH.LEFT),
        "DirectoryEntry2": ("Times New Roman", 9.5, False, WD_ALIGN_PARAGRAPH.LEFT),
        "DirectoryEntry3": ("Times New Roman", 9, False, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (font, size, bold, alignment) in custom.items():
        set_style(
            document,
            name,
            font,
            size,
            bold=bold,
            alignment=alignment,
        )
        style = document.styles[name]
        style.paragraph_format.space_after = Pt(6)
        if name.endswith("beschriftung"):
            style.paragraph_format.keep_with_next = True

    for level in range(1, 4):
        style = document.styles[f"DirectoryEntry{level}"]
        style.paragraph_format.left_indent = Cm((level - 1) * 0.5)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.tab_stops.add_tab_stop(
            Cm(16.0),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )

    set_style(document, "Source Code", "Consolas", 8.5)
    document.styles["Source Code"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.styles["Source Code"].paragraph_format.space_after = Pt(0)

    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Heading 1", "FrontHeading"}:
            set_bottom_border(paragraph, NAVY, 8)
        elif paragraph.style.name == "TitleRule":
            set_bottom_border(paragraph, RED, 10)
            paragraph.paragraph_format.space_after = Pt(16)

    for table_index, table in enumerate(document.tables):
        table.style = "Table"
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            row.height = None
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_shading(cell, LIGHT_GRAY if row_index == 0 else "FFFFFF")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(9.5)
                        if row_index == 0:
                            run.bold = True

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("DirectoryEntry"):
            target_text = paragraph.text.rsplit("\t", 1)[0].strip()
            target_text = re.sub(r"\s+\?\?$", "", target_text)
            page = page_map.get(target_text, "??") if page_map else "??"
            paragraph.text = f"{target_text}\t{page}"

    chapter_markers = {
        "1. Einleitung",
        "4. Implementierte Systemarchitektur",
        "6. Persistenz, Audit und Dokumente",
        "9. Risiken und technische Schulden",
    }
    found = set()
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() in chapter_markers:
            insert_section_break_before(paragraph)
            found.add(paragraph.text.strip())
    missing = chapter_markers - found
    if missing:
        raise RuntimeError(f"Kapitelanfänge nicht gefunden: {sorted(missing)}")

    intermediate = target.with_suffix(".sections.docx")
    document.save(intermediate)
    document = Document(intermediate)
    intermediate.unlink()

    if len(document.sections) < 5:
        raise RuntimeError(
            f"Erwartet wurden 5 Word-Abschnitte, gefunden: {len(document.sections)}"
        )

    authors = (
        "Nils Firschau",
        "Nils Firschau",
        "Paul Faller",
        "Robin Steiner",
        "Ole Schildt",
    )
    for index, section in enumerate(document.sections):
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.gutter = Cm(0)
        set_header(section, authors[min(index, len(authors) - 1)], index == 0)
        if index == 0:
            set_footer(section, fmt="upperRoman", start=0, blank_first=True)
        elif index == 1:
            set_footer(section, fmt="decimal", start=1)
        else:
            set_footer(section, fmt="decimal", start=None)

    document.core_properties.title = (
        "Modernisierung der Vereinsverwaltung FH_MA — Architekturkonzeption und Umsetzungsevaluation"
    )
    document.core_properties.subject = "Seminararbeit Advanced Software Engineering"
    document.core_properties.author = (
        "Nils Firschau; Paul Faller; Robin Steiner; Ole Schildt"
    )
    document.save(target)


def main() -> int:
    if len(sys.argv) not in {3, 4}:
        print(
            "Usage: format_docx.py <input.docx> <output.docx> [page-map.json]",
            file=sys.stderr,
        )
        return 2
    page_map = None
    if len(sys.argv) == 4:
        page_map = json.loads(Path(sys.argv[3]).read_text(encoding="utf-8"))
    format_document(Path(sys.argv[1]), Path(sys.argv[2]), page_map)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
